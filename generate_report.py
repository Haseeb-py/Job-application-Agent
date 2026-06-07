"""Generate the Job Application Agent technical project report as a DOCX file."""

from __future__ import annotations

from datetime import date
import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

logger = logging.getLogger(__name__)

REPORT_PATH = Path("Job_Application_Agent_Project_Report.docx")


def _add_title_page(document: Document) -> None:
    """Add the report title page.

    Args:
        document: DOCX document being generated.

    Returns:
        None.
    """

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Job Application Agent — Technical Project Report")
    run.bold = True
    run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Autonomous AI-Powered Job Search & Application System")
    sub_run.italic = True
    sub_run.font.size = Pt(14)

    document.add_paragraph()
    fields = document.add_paragraph()
    fields.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fields.add_run("Author: ______________________________\n").bold = True
    fields.add_run("University: ___________________________\n").bold = True
    fields.add_run(f"Date: {date.today().isoformat()}").bold = True
    document.add_page_break()


def _add_heading(document: Document, title: str, level: int = 1) -> None:
    """Add a section heading.

    Args:
        document: DOCX document being generated.
        title: Heading title.
        level: Heading level.

    Returns:
        None.
    """

    document.add_heading(title, level=level)


def _add_paragraph(document: Document, text: str) -> None:
    """Add a normal paragraph.

    Args:
        document: DOCX document being generated.
        text: Paragraph text.

    Returns:
        None.
    """

    document.add_paragraph(text)


def _add_bullets(document: Document, items: list[str]) -> None:
    """Add bullet points.

    Args:
        document: DOCX document being generated.
        items: Bullet text items.

    Returns:
        None.
    """

    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Add a table to the document.

    Args:
        document: DOCX document being generated.
        headers: Header labels.
        rows: Table rows.

    Returns:
        None.
    """

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def build_report() -> Document:
    """Build the complete technical report document.

    Args:
        None.

    Returns:
        Populated DOCX Document instance.
    """

    document = Document()
    _add_title_page(document)

    _add_heading(document, "Abstract")
    _add_paragraph(
        document,
        "The Job Application Agent is an autonomous AI-powered system that assists candidates in discovering relevant jobs, "
        "matching those jobs against a resume, generating tailored cover letters, analyzing skill gaps, and exporting a structured "
        "application payload. The project combines web automation, natural language processing, vector similarity search, workflow "
        "orchestration, and an interactive Streamlit interface. It is designed as a practical final-year Computer Science or Data "
        "Science project because it demonstrates end-to-end integration of modern AI engineering concepts rather than a single isolated model.",
    )

    _add_heading(document, "Project Objectives")
    _add_bullets(
        document,
        [
            "Automate job discovery from LinkedIn and public Glassdoor listings using Selenium-based scraping.",
            "Parse PDF and DOCX resumes and extract technical skills from a fixed professional taxonomy.",
            "Rank jobs by semantic relevance using sentence-transformers embeddings and FAISS vector search.",
            "Generate professional three-paragraph cover letters with LangChain and the Groq Llama 3 model.",
            "Export application-ready JSON records containing job details, relevance scores, skill gaps, and cover letters.",
            "Provide a dark-themed Streamlit dashboard that makes the entire workflow usable by a non-technical candidate.",
        ],
    )

    _add_heading(document, "System Architecture")
    _add_paragraph(
        document,
        "The system follows a staged pipeline controlled by LangGraph. The user uploads a resume and configuration through Streamlit. "
        "The graph parses the resume, scrapes job listings, performs semantic matching, generates cover letters, and exports the final JSON. "
        "Each stage writes to a shared AgentState object, which makes the pipeline transparent and easy to debug.",
    )
    _add_paragraph(
        document,
        "Pipeline: Streamlit UI -> Resume Parser -> LinkedIn and Glassdoor Scrapers -> Skill Extraction -> FAISS Semantic Matcher -> "
        "Groq Cover Letter Generator -> Skill Gap Analyzer -> JSON Exporter.",
    )

    _add_heading(document, "Technology Stack")
    _add_table(
        document,
        ["Component", "Technology", "Purpose"],
        [
            ["Workflow orchestration", "LangGraph", "Represents the agent as a state machine with deterministic transitions."],
            ["LLM integration", "LangChain + Groq llama3-8b-8192", "Generates tailored cover letters from few-shot prompts."],
            ["Embeddings", "sentence-transformers all-MiniLM-L6-v2", "Converts resumes and jobs into semantic vectors."],
            ["Vector search", "FAISS IndexFlatL2", "Finds the most relevant job descriptions efficiently."],
            ["Scraping", "Selenium, undetected-chromedriver, BeautifulSoup4", "Renders and extracts job listing content."],
            ["Resume parsing", "PyPDF2, python-docx", "Extracts text from PDF and DOCX resumes."],
            ["Interface", "Streamlit", "Provides multi-tab user interaction and visualization."],
            ["Reliability", "tenacity", "Retries Groq LLM calls with exponential backoff."],
        ],
    )

    _add_heading(document, "LangGraph Agent Design")
    _add_paragraph(
        document,
        "The AgentState TypedDict stores resume text, extracted skills, raw listings, matched jobs, cover letters, final output records, errors, "
        "and the current pipeline status. LangGraph nodes are implemented as pure state-transforming functions. This design makes the workflow "
        "modular: each node can be tested independently, and failures can be routed to an error handler without mixing UI code with business logic.",
    )
    _add_bullets(
        document,
        [
            "scrape_jobs_node collects LinkedIn and Glassdoor jobs and deduplicates them by URL.",
            "match_jobs_node extracts resume skills and filters jobs by semantic relevance.",
            "generate_cover_letters_node performs skill-gap analysis and calls Groq with retry/backoff.",
            "export_output_node assembles the final JSON application records.",
            "error_handler_node logs accumulated errors and marks the workflow as failed.",
        ],
    )

    _add_heading(document, "Scraping Strategy")
    _add_paragraph(
        document,
        "LinkedIn scraping uses credential-based login because the platform usually hides meaningful job details behind authentication. "
        "The scraper loads credentials from environment variables, uses undetected Chrome in headless mode, waits for job cards, clicks each card, "
        "and extracts title, company, location, description, URL, and posting date. All Selenium operations are wrapped with try/except handling.",
    )
    _add_paragraph(
        document,
        "Glassdoor scraping intentionally avoids login. It opens public job search pages, dismisses cookie banners and modal pop-ups when possible, "
        "then parses the rendered DOM with BeautifulSoup. Publicly unavailable fields such as salary are set to None instead of causing a crash.",
    )

    _add_heading(document, "Resume Parsing and Skill Extraction")
    _add_paragraph(
        document,
        "The parser supports PDF and DOCX resumes. PyPDF2 extracts text from each PDF page, while python-docx reads paragraph text from DOCX files. "
        "The extracted text is returned with a word count. Skill extraction uses a hardcoded taxonomy of more than eighty technologies and performs "
        "case-insensitive exact-term matching to avoid noisy substring matches.",
    )

    _add_heading(document, "Semantic Matching Method")
    _add_paragraph(
        document,
        "The matcher loads the all-MiniLM-L6-v2 sentence-transformers model once at module import time. Job descriptions are embedded into vectors, "
        "normalized, and indexed in a FAISS IndexFlatL2 index. The resume is encoded as a single vector and searched against the index. Distances are "
        "converted into normalized relevance scores between 0 and 1, filtered by threshold, and sorted in descending order.",
    )

    _add_heading(document, "Cover Letter Generation")
    _add_paragraph(
        document,
        "Cover letters are generated using a LangChain PromptTemplate with two complete few-shot examples: one for a Software Engineer role and one "
        "for a Data Scientist role. The prompt instructs the model to write exactly three professional paragraphs, avoid fluff, highlight matching "
        "skills, and show enthusiasm for the specific company. Groq calls are protected with tenacity retry logic to handle transient API failures.",
    )

    _add_heading(document, "Output Schema")
    _add_paragraph(
        document,
        "Each exported application record contains job_id, job_title, company, location, URL, source, relevance_score, matching_skills, "
        "missing_skills, match_percentage, cover_letter, and scraped_at. The exporter creates parent directories automatically and writes indented JSON.",
    )

    _add_heading(document, "Streamlit User Interface")
    _add_paragraph(
        document,
        "The Streamlit application provides a dark-themed dashboard with sidebar configuration. Users upload a resume, choose job title and location, "
        "set maximum results and threshold, then run the agent. Results are presented across four tabs: Matched Jobs, Cover Letters, Skill Gap Analysis, "
        "and Export. This interface makes the system demonstrable in a viva because every internal stage has a visible output.",
    )

    _add_heading(document, "Testing and Validation Strategy")
    _add_bullets(
        document,
        [
            "Unit test resume parsing with small sample PDF and DOCX files.",
            "Unit test skill extraction with mixed-case skill names and boundary-sensitive terms such as C++ and Node.js.",
            "Mock scraper outputs to validate LangGraph transitions without relying on live websites.",
            "Mock Groq responses to validate cover letter and export assembly.",
            "Run manual end-to-end validation through Streamlit with a real resume and controlled job search query.",
        ],
    )

    _add_heading(document, "Security and Ethics")
    _add_paragraph(
        document,
        "The project avoids hardcoded credentials by loading secrets from environment variables. The example environment file uses placeholders. "
        "Scraping should be performed responsibly, at low volume, and in accordance with platform terms. The system does not submit applications automatically; "
        "it prepares application materials for human review, which keeps the candidate in control.",
    )

    _add_heading(document, "Known Limitations")
    _add_bullets(
        document,
        [
            "LinkedIn may block automation with verification checkpoints or bot-detection systems.",
            "Glassdoor public pages expose limited information without login, so salary and full descriptions may be unavailable.",
            "Semantic relevance can be weak when a job description is only a short snippet.",
            "The skill taxonomy is fixed, so new technologies must be added manually unless a dynamic extraction model is introduced.",
            "The project requires Chrome, network access, and valid API credentials for full operation.",
        ],
    )

    _add_heading(document, "Viva and Interview Preparation")
    _add_paragraph(
        document,
        "A strong explanation of this project should emphasize why LangGraph was used for controllable orchestration, why FAISS was selected for "
        "fast vector similarity search, and why the system separates scraping, parsing, matching, generation, and export into independent modules. "
        "The candidate should also be ready to discuss scraping risks, credential security, model hallucination, prompt design, retry logic, and how "
        "the system could be extended with application tracking, email drafting, or human approval workflows.",
    )

    _add_heading(document, "Conclusion")
    _add_paragraph(
        document,
        "The Job Application Agent demonstrates a complete applied AI workflow that combines data acquisition, NLP, vector search, agent orchestration, "
        "LLM generation, and user-facing software engineering. It is not merely a chatbot; it is a structured automation system with clear state, "
        "error handling, typed modules, and exportable results. These qualities make it suitable for academic demonstration and practical portfolio use.",
    )

    return document


def main() -> None:
    """Generate and save the DOCX report.

    Args:
        None.

    Returns:
        None.
    """

    logging.basicConfig(level=logging.INFO)
    document = build_report()
    document.save(REPORT_PATH)
    logger.info("Saved report to %s", REPORT_PATH)


if __name__ == "__main__":
    main()

